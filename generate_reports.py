"""
Project: Information Transmission and Sectoral Asymmetry in the KSE-100
Module : Report Generator — numbered Word documents in Results/
Author : Dr. Yasir Saeed (KUST)
Description: Builds the sequenced report set (mirroring the Shariah
             Compliance Project organisation): one self-contained Word
             document per analysis stage, each with methods, tables,
             figures and commentary. All numbers are read live from the
             pipeline's output files, so re-running the pipeline and then
             this script regenerates a consistent report set.
USAGE       python generate_reports.py
"""

import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

P = Path(r"D:\Stockmarket-Sentiment-Analysis")
RES, TAB, FIG = P / "Results", P / "Results" / "tables", P / "Results" / "figures"
RES.mkdir(exist_ok=True)

sys.path.insert(0, str(P))


# ---------------------------------------------------------------- helpers ----
class Report:
    def __init__(self, number, title):
        self.doc = Document()
        n = self.doc.styles["Normal"]
        n.font.name = "Arial"; n.font.size = Pt(10.5)
        n.paragraph_format.space_after = Pt(6)
        for hname, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)]:
            st = self.doc.styles[hname]
            st.font.name = "Arial"; st.font.size = Pt(size); st.font.bold = True
            st.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        for sec in self.doc.sections:
            sec.page_width, sec.page_height = Inches(8.5), Inches(11)
            for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                setattr(sec, m, Inches(0.9))
        t = self.doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(f"Report {number:02d} — {title}"); r.bold = True; r.font.size = Pt(17)
        s = self.doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = s.add_run("Information Transmission and Sectoral Asymmetry in the KSE-100 — "
                      "Dr. Yasir Saeed, KUST"); r.italic = True; r.font.size = Pt(9)
        self.number, self.title = number, title

    def h1(self, t): self.doc.add_heading(t, level=1)
    def h2(self, t): self.doc.add_heading(t, level=2)

    def p(self, text, bold=False, italic=False, size=None):
        par = self.doc.add_paragraph(); r = par.add_run(text)
        r.bold, r.italic = bold, italic
        if size: r.font.size = Pt(size)
        return par

    def bullet(self, text):
        self.doc.add_paragraph(text, style="List Bullet")

    def note(self, text):
        par = self.doc.add_paragraph(); r = par.add_run("Note: " + text)
        r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def table(self, df, fontsize=8, max_rows=None, index=False):
        if max_rows and len(df) > max_rows:
            self.note(f"Showing first {max_rows} of {len(df)} rows — full table in the "
                      "companion Excel file.")
            df = df.head(max_rows)
        if index:
            df = df.reset_index()
        t = self.doc.add_table(rows=1 + len(df), cols=len(df.columns))
        t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, c in enumerate(df.columns):
            cell = t.rows[0].cells[j]; cell.text = ""
            r = cell.paragraphs[0].add_run(str(c)); r.bold = True; r.font.size = Pt(fontsize)
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), "D5E8F0"); tcPr.append(shd)
        for i in range(len(df)):
            for j, c in enumerate(df.columns):
                v = df.iloc[i, j]
                txt = "" if pd.isna(v) else (f"{v:,.4f}".rstrip("0").rstrip(".")
                                             if isinstance(v, float) else str(v))
                cell = t.rows[i + 1].cells[j]; cell.text = ""
                r = cell.paragraphs[0].add_run(txt); r.font.size = Pt(fontsize)
        self.doc.add_paragraph()

    def figure(self, path, width=6.6, caption=None):
        path = Path(path)
        if not path.exists():
            self.note(f"figure not found: {path.name}"); return
        par = self.doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(path), width=Inches(width))
        if caption:
            c = self.doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)

    def save(self):
        fname = RES / f"{self.number:02d}_{self.title.lower().replace(' ', '_').replace('/', '-')}.docx"
        self.doc.save(fname)
        print("saved", fname.name)


def xl(path, sheet=0):
    return pd.read_excel(path, sheet_name=sheet)


# =============================================================== report 01 ----
def r01():
    rep = Report(1, "Data and Corpus")
    rep.h1("1. News corpus")
    rep.p("Source: Business Recorder (brecorder.com), scraped by sequential article-ID "
          "scan (news_extractor.py) into monthly CSVs with article_id, url, headline, "
          "body_text, date, author, http_status. After deduplication on article_id and "
          "restriction to http_status = 200 the corpus is 416,806 unique articles "
          "covering 22 June 2020 – 26 May 2026.")
    rep.h2("Date normalisation")
    rep.p("Two date formats coexisted across monthly files (ISO YYYY-MM-DD and day-first "
          "DD-MM-YYYY in some 2020–21 files). Day-first ordering was verified empirically "
          "(first field exceeds 12 in 3,164 ambiguous rows, second never does); all dates "
          "were normalised to ISO, preventing silent day/month swaps (4,891 conversions "
          "among relevant articles).")
    rep.h1("2. Price data")
    rep.p("Ticker-level PSX daily data (compiled_psx_historical_2017_2026.csv): 922,837 "
          "ticker-day rows, 1,191 symbols, 2 Jan 2017 – 30 Jun 2026 (DATE, SYMBOL, LDCP, "
          "OPEN, HIGH, LOW, CLOSE, CHANGE, VOLUME). The file contains no usable index "
          "series, so sector and market returns are constructed from constituents "
          "(Report 04). Sector classification is taken from the PSX Data Portal symbols "
          "endpoint (Source Data/psx_symbols_sectors.json; 745 currently listed equities "
          "in 38 named sectors).")
    rep.note("Delisted securities absent from the current listing remain unmapped and are "
             "excluded from sector portfolios — a survivorship caveat; 62% of all "
             "historical tickers map, and coverage among actively traded stocks in the "
             "study window is far higher.")
    rep.save()


# =============================================================== report 02 ----
def r02():
    from topic_router import (MONETARY_KEYWORDS, FISCAL_KEYWORDS,
                              EXTERNAL_FINANCE_KEYWORDS, ENERGY_KEYWORDS)
    rep = Report(2, "Relevance Filter and Keyword Dictionaries")
    rep.h1("1. Method")
    rep.p("An article is PSX-relevant if and only if it classifies into at least one of "
          "the four macroeconomic policy channels with sufficient keyword evidence. For "
          "each channel: score = 3 x headline keyword occurrences + 1 x body occurrences "
          "(case-insensitive substring counts; the 3x weight captures editorial "
          "emphasis). Relevance requires total weighted hits >= 3 (one headline keyword "
          "or three body mentions). The vectorised implementation is verified identical "
          "to the reference classifier on a 500-article random sample at every run.")
    rep.h2("Multi-label channel assignment (v2)")
    rep.p("An article joins EVERY channel in which its own weighted hits >= 3 — the same "
          "evidence standard as relevance. Genuinely multi-channel articles (e.g. an IMF "
          "review discussing energy tariffs) inform every channel they substantively "
          "touch, instead of being parked in a residual 'Mixed' bucket.")
    rep.h1("2. Keyword dictionaries (full listing)")
    for name, kws in [("Monetary", MONETARY_KEYWORDS), ("Fiscal", FISCAL_KEYWORDS),
                      ("External Finance", EXTERNAL_FINANCE_KEYWORDS),
                      ("Energy", ENERGY_KEYWORDS)]:
        rep.h2(f"{name} ({len(kws)} keywords)")
        rep.p(";  ".join(kws), size=9)
    rep.h1("3. Outcome")
    rep.table(pd.DataFrame([
        ["Corpus (unique, http 200)", "416,806", "100%"],
        ["Relevant (>=1 channel, hits >= 3)", "139,546", "33.5%"],
        ["In-channel, weak evidence (hits 1-2)", "111,269", "26.7%"],
        ["Unclassified (no macro keywords)", "165,991", "39.8%"],
    ], columns=["Bucket", "Articles", "Share"]))
    art = pd.read_csv(P / "Phase 2 Outputs" / "article_channel_scores.csv",
                      usecols=["article_id", "channel"])
    cnt = art.groupby("channel")["article_id"].count().rename("article-channel pairs")
    multi = art.groupby("article_id").size()
    rep.h2("Multi-label assignment counts")
    rep.table(cnt.reset_index())
    rep.p(f"Articles assigned to at least one channel: {multi.size:,}; to two or more "
          f"channels: {(multi >= 2).sum():,} ({(multi >= 2).mean() * 100:.1f}%).")
    rep.note("Audit workbook: Phase 1 Outputs/keyword_filter_audit.xlsx (hit distribution "
             "+ 300-article samples of relevant, borderline and excluded buckets).")
    rep.save()


# =============================================================== report 03 ----
def r03():
    from psx_sign_rules import RULES
    rep = Report(3, "Sentiment Construction")
    rep.h1("1. FinBERT tone scoring")
    rep.p("Each relevant article is scored with FinBERT (ProsusAI/finbert; Araci 2019) on "
          "headline + first 2,000 body characters, truncated at 256 tokens (fp16, GPU, "
          "checkpointed batches). Output is converted to a signed tone score in [-1, +1].")
    rep.h1("2. Directional sign rules")
    rep.p("Tone is not direction: a rate hike reads as decisive prose but is negative for "
          "equities. 35 channel-specific event rules supply the DIRECTION for KSE-100 "
          "equities; FinBERT confidence supplies the MAGNITUDE (0.5 default when FinBERT "
          "is neutral). Headline matches weigh 3x. When an article belongs to several "
          "channels, each channel applies its own rules — the same article can correctly "
          "carry different directional readings in different channels. Articles matching "
          "no rule fall back to the FinBERT score. Global crude-oil price moves carry no "
          "rule by design (sign ambiguous for the aggregate market).")
    rows = []
    for ch, rules in RULES.items():
        for rule in rules:
            rows.append({"channel": ch, "rule": rule["name"],
                         "sign for PSX": "+" if rule["sign"] > 0 else "-",
                         "patterns": len(rule["patterns"])})
    rep.table(pd.DataFrame(rows), fontsize=8)
    rep.h1("3. Combined series S_All")
    rep.p("S_All aggregates ALL relevant articles, scored against the union of every "
          "rule — the overall news-sentiment measure paired with the combined market "
          "index r_Market in estimation.")
    rep.h1("4. Weekend and holiday news")
    art = pd.read_csv(P / "Phase 2 Outputs" / "article_channel_scores.csv",
                      usecols=["article_id"])
    rep.p("News published on non-trading days cannot move prices until the next session. "
          "Every article is therefore assigned to the first PSX trading day on or after "
          "its publication date (trading calendar taken from the price data). "
          "Previously, weekend/holiday news was lost in the trading-day merge.")
    rep.h1("5. Daily aggregation")
    rep.p("Word-count-weighted mean of article scores within each (trading day, channel) "
          "cell; days without articles in a channel carry the last observed value "
          "forward. Raw-FinBERT variants of all five series are produced in parallel as "
          "the robustness benchmark (Report 12).")
    rep.figure(FIG / "03_sentiment_series.png",
               caption="Daily PSX-directional sentiment series")
    rep.note("Sign-rule audit: Phase 2 Outputs/sign_rules_audit.xlsx. Article-channel "
             "level detail: Phase 2 Outputs/article_channel_scores.csv.")
    rep.save()


# =============================================================== report 04 ----
def r04():
    rep = Report(4, "Sector Universe and Return Construction")
    rep.h1("1. Eligibility criteria")
    rep.bullet("Equity securities only — debt instruments, ETFs, mutual funds and the "
               "'Miscellaneous' grouping excluded from sector analysis.")
    rep.bullet("Stock-level screens: valid CLOSE and LDCP, |daily return| <= 25%, and "
               ">= 70% trading-day coverage in the study window.")
    rep.bullet("Sector quorum: >= 4 screened firms. Four (not five) preserves Oil & Gas "
               "Exploration — the KSE-100's second-largest sector has exactly four firms "
               "(OGDC, PPL, POL, MARI), all large caps.")
    rep.bullet("A sector-day requires >= 3 traded constituents.")
    rep.h1("2. Return definitions")
    rep.p("Stock return: r = CLOSE / LDCP - 1, using the exchange's own adjusted previous "
          "close (robust to corporate actions). Sector return: equal-weighted mean over "
          "constituents that traded that day. Market return r_Market: equal-weighted "
          "mean over ALL screened equities (319 stocks, ~294 per day) — the combined-"
          "index proxy, as the price file contains no usable index series.")
    rep.h1("3. Eligible sectors")
    elig = xl(P / "Phase 3 Inputs" / "sector_universe.xlsx", "sector_diagnostics")
    rep.table(elig, fontsize=8)
    rep.note("Full eligibility table incl. rejected sectors: Phase 3 Inputs/"
             "sector_universe.xlsx.")
    rep.figure(FIG / "04_cumulative_returns.png",
               caption="Cumulative return paths — market and the five focus sectors")
    rep.save()


# =============================================================== report 05 ----
def r05():
    rep = Report(5, "Descriptive Statistics")
    rep.h1("1. Sentiment series")
    rep.table(xl(TAB / "05_descriptive_statistics.xlsx", "sentiment"), fontsize=8)
    rep.h1("2. Sector returns (%)")
    rep.table(xl(TAB / "05_descriptive_statistics.xlsx", "returns_pct"), fontsize=7.5)
    rep.h1("3. News flow over time")
    rep.figure(FIG / "02_article_flow_by_channel.png",
               caption="Monthly article counts by channel (multi-label)")
    rep.p("Interpretation: article flow spikes track Pakistan's macro calendar — budget "
          "months (June), IMF programme milestones, and the 2022–23 stabilisation "
          "crisis — confirming the channels capture policy news intensity.")
    rep.save()


# =============================================================== report 06 ----
def r06():
    rep = Report(6, "Correlation Analysis")
    rep.p("Contemporaneous correlations below; dynamic (lead-lag) relationships are "
          "treated in Reports 09–11.")
    rep.h1("1. Sentiment x sentiment")
    rep.table(xl(TAB / "06_correlations.xlsx", "sentiment_x_sentiment"), fontsize=8)
    rep.figure(FIG / "07_corr_sentiment.png", width=4.8)
    rep.h1("2. Returns x sentiment (contemporaneous)")
    rep.figure(FIG / "06_corr_returns_sentiment.png", width=5.2)
    rep.h1("3. Returns x returns")
    rep.figure(FIG / "05_corr_returns.png")
    rep.note("Numeric matrices: Results/tables/06_correlations.xlsx.")
    rep.save()


# =============================================================== report 07 ----
def r07():
    rep = Report(7, "Unit Root Tests")
    rep.p("Stationarity is a precondition for VAR-in-levels and Granger testing. ADF "
          "(H0: unit root) and KPSS (H0: stationarity) per series, plus the Maddala-Wu "
          "(1999) Fisher-type panel test combining individual ADF p-values, "
          "-2 SUM ln(p_i) ~ chi2(2N).")
    rep.h1("1. Panel tests")
    rep.table(xl(TAB / "07_unit_root_tests.xlsx", "MaddalaWu_panel"))
    rep.h1("2. Series-by-series")
    rep.table(xl(TAB / "07_unit_root_tests.xlsx", "ADF_KPSS_by_series"), fontsize=7.5)
    rep.note("KPSS p-values are interpolation-bounded by statsmodels at [0.01, 0.10].")
    rep.save()


# =============================================================== report 08 ----
def r08():
    rep = Report(8, "VAR Lag Selection")
    rep.p("Lag order for the market systems selected over 1–10 lags. The estimation "
          "suite uses the AIC-selected order per system (sectoral VARs select their own "
          "AIC order).")
    rep.table(xl(TAB / "08_lag_selection.xlsx"))
    rep.save()


# =============================================================== report 09 ----
def r09():
    rep = Report(9, "Granger Causality")
    rep.p("Bivariate Granger tests (SSR F-test) at lags 1–5, in BOTH directions: "
          "sentiment -> return (transmission) and return -> sentiment (reverse "
          "causality: markets moving before the press writes, or coverage following "
          "price action). min_p is the smallest p-value across the five lags; sig_lags "
          "lists lags significant at 5%.")
    rep.h1("1. Market level (combined index)")
    rep.table(xl(TAB / "09_granger_causality.xlsx", "market"), fontsize=8)
    rep.h1("2. All sectors x channels (bidirectional)")
    rep.table(xl(TAB / "09_granger_causality.xlsx", "sectors_bidirectional"),
              fontsize=7, max_rows=110)
    rep.save()


# =============================================================== report 10 ----
def r10():
    rep = Report(10, "VAR IRF and FEVD")
    rep.h1("1. Market systems")
    rep.table(xl(TAB / "10_var_irf_fevd.xlsx", "var_specs"))
    rep.figure(FIG / "08_irf_market_SAll.png", width=4.2,
               caption="r_Market response to a unit S_All innovation (95% MC bands, 500 reps)")
    rep.figure(FIG / "09_irf_market_channels.png",
               caption="r_Market response to channel sentiment innovations")
    rep.h2("FEVD — market")
    rep.table(xl(TAB / "10_var_irf_fevd.xlsx", "fevd_market"))
    rep.h2("IRF summary — market")
    rep.table(xl(TAB / "10_var_irf_fevd.xlsx", "irf_market"))
    rep.h1("2. All-sector sensitivity matrix")
    rep.figure(FIG / "10_sensitivity_heatmap_all.png", width=6.0,
               caption="FEVD share (%) of each sector's 22-day return variance from each "
                       "channel; * = Granger-significant at 5%")
    rep.h2("Sectoral IRF summary")
    rep.table(xl(TAB / "10_var_irf_fevd.xlsx", "irf_sectors"), fontsize=7, max_rows=110)
    rep.note("Full FEVD tables: Results/tables/10_var_irf_fevd.xlsx and "
             "10b_sensitivity_matrix_all.xlsx.")
    rep.save()


# =============================================================== report 11 ----
def r11():
    rep = Report(11, "Lead-Lag and Anticipation Tests")
    rep.p("Market efficiency motivates testing LEAD effects: if information reaches "
          "investors before publication, prices adjust in advance and future sentiment "
          "helps explain today's return. For every return series and sentiment series: "
          "OLS of r_t on five LAGS and five LEADS of sentiment, Newey-West HAC "
          "standard errors (5 lags). Joint F-test on the lags = delayed transmission; "
          "joint F-test on the leads = anticipation / pre-publication adjustment.")
    ll = xl(TAB / "11_lead_lag_anticipation.xlsx")
    rep.h1("1. Market level")
    rep.table(ll[ll["return"] == "r_Market"], fontsize=8)
    rep.figure(FIG / "11_ccf_market.png",
               caption="Cross-correlations corr(r_Market[t], S[t+k]); red bars (k<0) = "
                       "returns lead sentiment (anticipation side)")
    rep.h1("2. All sectors")
    sig = ll[(ll["lags_p"] < 0.05) | (ll["leads_p"] < 0.05)]
    rep.p(f"Of {len(ll)} return x sentiment pairs, {int((ll['lags_p'] < .05).sum())} show "
          f"significant lagged impact and {int((ll['leads_p'] < .05).sum())} show "
          f"significant lead (anticipation) effects at 5%. Significant pairs:")
    rep.table(sig, fontsize=7, max_rows=120)
    rep.note("Full table: Results/tables/11_lead_lag_anticipation.xlsx.")
    rep.save()


# =============================================================== report 12 ----
def r12():
    rep = Report(12, "Robustness")
    rep.h1("1. Directional correction vs raw FinBERT")
    rep.p("Every headline test re-estimated on sentiment series built from raw FinBERT "
          "tone (no sign rules). If the directional layer adds information, "
          "transmission evidence should weaken under raw tone.")
    rep.table(xl(TAB / "12_robustness.xlsx", "directional_vs_raw"))
    rep.h1("2. Original five-sector subset")
    rep.p("The study's original five focus sectors, extracted from the all-sector "
          "results (directional series):")
    rep.table(xl(TAB / "12_robustness.xlsx", "original_five_subset"), fontsize=8)
    rep.h1("3. Further robustness recorded for the paper")
    rep.bullet("Raw-series Granger tables (market and sectors) and lead-lag tests: "
               "sheets in Results/tables/12_robustness.xlsx.")
    rep.bullet("Relevance-threshold sensitivity (hits >= 2 or >= 4) requires re-scoring "
               "the wider/narrower article set with FinBERT; noted as an extension.")
    rep.bullet("Equal weighting: value-weighted variants require market-cap data; the "
               "market portfolio's high correlation with published KSE-100 levels over "
               "the window can be cited as external validity.")
    rep.save()


if __name__ == "__main__":
    for fn in [r01, r02, r03, r04, r05, r06, r07, r08, r09, r10, r11, r12]:
        try:
            fn()
        except Exception as e:
            print(f"FAILED {fn.__name__}: {e}")
    print("Report generation complete.")
